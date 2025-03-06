#!/usr/bin/env python3
import os
import re
import glob
import yaml
import argparse
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.shared import Pt

def set_arial_font(run):
    """Set the font to Arial for a run object."""
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run._element.rPr.rFonts.set(qn('w:cs'), 'Arial')
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')

def extract_frontmatter(content):
    """
    Extracts and returns the YAML front matter as a dict,
    and also returns the content after the front matter.
    If no valid front matter is found, returns (None, content).
    """
    frontmatter_regex = re.compile(r"^---\s*\n(.*?)\n---\s*\n", re.DOTALL)
    match = frontmatter_regex.match(content)
    if match:
        frontmatter_str = match.group(1)
        try:
            data = yaml.safe_load(frontmatter_str)
        except yaml.YAMLError as e:
            print(f"Error parsing YAML: {e}")
            data = None
        return data, content[match.end():]
    return None, content

def md_to_docx(input_dir, output_file):
    doc = Document()
    md_file_paths = glob.glob(os.path.join(input_dir, "*.md"))
    files_with_order = []

    # Extract order from YAML front matter in each file
    for filepath in md_file_paths:
        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read()
        frontmatter, remaining = extract_frontmatter(content)
        order = frontmatter.get("order") if (frontmatter and "order" in frontmatter) else float('inf')
        title = frontmatter.get("title") if frontmatter else None
        group = frontmatter.get("group") if frontmatter else None
        files_with_order.append((order, filepath, remaining, title, group))

    # Sort files based on the order value
    files_with_order.sort(key=lambda tup: tup[0])

    current_group = None

    for idx, (_, filepath, content, title, group) in enumerate(files_with_order):
        print(f"Processing {filepath} ...")

        # Add page break before content (except for first file)
        if idx > 0:
            doc.add_page_break()

        # Add group as header 1 only if it's a new group
        if group and group != current_group:
            heading = doc.add_heading(group, level=1)
            for run in heading.runs:
                set_arial_font(run)
            current_group = group

        # Add title as header 2 if it exists
        if title:
            heading = doc.add_heading(title, level=2)
            for run in heading.runs:
                set_arial_font(run)

        # Process content line by line
        for line in content.splitlines():
            # Check for markdown header syntax
            header_match = re.match(r"^(#{1,6})\s+(.*)$", line)
            if header_match:
                hashes, text = header_match.groups()
                level = len(hashes) + 1  # Demote header level by 1
                if level > 6:  # Cap at level 6
                    level = 6
                heading = doc.add_heading(text.strip(), level=level)
                for run in heading.runs:
                    set_arial_font(run)
            else:
                # Add non-header text lines as paragraphs
                paragraph = doc.add_paragraph(line.rstrip())
                for run in paragraph.runs:
                    set_arial_font(run)

    doc.save(output_file)
    print(f"Saved DOCX output to {output_file}")

if __name__ == '__main__':
    parser = argparse.ArgumentParser(
        description="Combine Markdown files (ordered by eleventy front matter) into a single DOCX document"
    )
    parser.add_argument(
        "input_dir",
        help="Directory containing Markdown (.md) files."
    )
    parser.add_argument(
        "output_file",
        help="Path to the output DOCX file (e.g. output.docx)."
    )
    args = parser.parse_args()
    md_to_docx(args.input_dir, args.output_file)
